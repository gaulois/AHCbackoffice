import io
import tempfile
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def add_qr_to_word(input_file_stream, qr_bytes, client_id):
    """
    Insère un QR code dans un tableau à 2 colonnes du pied de page :
    - colonne gauche : QR code
    - colonne droite : contenu existant du pied de page

    Renvoie un flux binaire prêt à être uploadé dans MinIO.
    """
    # Sauvegarde temporaire du fichier Word original
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        tmp_file.write(input_file_stream.read())
        tmp_file_path = tmp_file.name

    # Ouvre le document
    doc = Document(tmp_file_path)

    # Utilise la première section
    section = doc.sections[0]
    footer = section.footer

    # Récupère le texte existant du footer
    existing_text_lines = []
    for p in footer.paragraphs:
        txt = p.text.strip()
        if txt:
            existing_text_lines.append(txt)

    existing_text = "\n".join(existing_text_lines)

    # Supprime le contenu existant du footer
    # (python-docx ne propose pas de clear() direct, donc on vide les paragraphes)
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    # Crée un tableau 1 ligne / 2 colonnes dans le footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    row = table.rows[0]
    left_cell = row.cells[0]
    right_cell = row.cells[1]

    # Largeurs visuelles approximatives
    left_cell.width = Inches(1.2)
    right_cell.width = Inches(5.3)

    # Crée une image temporaire du QR
    qr_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    qr_temp.write(qr_bytes.getvalue())
    qr_temp.close()

    # Colonne gauche : QR
    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_paragraph.add_run()
    left_run.add_picture(qr_temp.name, width=Inches(0.9))

    # Colonne droite : texte existant
    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if existing_text:
        right_run = right_paragraph.add_run(existing_text)
        right_run.font.size = Pt(9)

    # Sauvegarde du document modifié
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)

    return output_stream